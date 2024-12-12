import gradio as gr
import PyPDF2
import openai
from docx import Document

client = openai.OpenAI(api_key='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')



# Fonction pour formater le texte de GPT et éviter les symboles Markdown (#, *)
def format_gpt_response_to_word(doc, gpt_response):
    lines = gpt_response.split('\n')
    
    for line in lines:
        # Convertir les titres Markdown (par exemple, ### Titre) en Heading
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)

        # Convertir les puces Markdown (*) en liste à puces
        elif line.startswith('- **'):
            doc.add_paragraph(line.replace('**', '').strip(), style='ListBullet')
        
        # Ajouter des paragraphes standards pour le texte normal
        else:
            doc.add_paragraph(line.strip())


# Fonction pour générer le document Word avec la réponse de l'API GPT
def generate_word_document(gpt_response):
    doc = Document()

    doc.add_heading("Rapport d'analyse de la facture d'hydro Quebec", 0)

    # Ajouter le contenu GPT au document en le formatant correctement
    format_gpt_response_to_word(doc, gpt_response)

    # Enregistrer le document
    file_name = "rapport_analyse_rempli.docx"
    doc.save(file_name)
    return file_name

# Fonction principale pour analyser et générer le document Word
def analyze_and_generate(pdf_files, scenario_text):
    for pdf in pdf_files:
        reader = PyPDF2.PdfReader(pdf)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    full_prompt = f"""En se basant sur les informations ci-dessous, completer les parties (a remplir). La methodologie de calcul pour certains est mentionnee plus bas. Pour les calculs, donner juste le resultat final, ne montrez pas les etapes de calculs
            Informations nécessaires :
            • Nom du client : (à remplir)
            • Secteur d'activité : (à remplir)
            • Quantité et puissance des bornes déjà installées (si applicable) : (à remplir)
            • Quantité et puissance des bornes à installer dans le futur (si connu) : (à remplir)

            Résultat de l'analyse préliminaire :

            • Tarif General actuel : (à remplir)
            • Puissance max sans appel de puissance facturé : (à remplir)
            • Puissance réelle : (à remplir)               Mois : (à remplir)
            • Puissance facturée (kW): (à remplir)
            • Consommation max facturée : (à remplir)      Mois : (à remplir)

            Scenario : {scenario_text}
            • Tarif : (à remplir)
            • Puissance à facturer selon le projet : (à remplir)
            • Coût mensuel estimatif (puissance) : (à remplir)
            • Impact mensuel de l’ajout des bornes : (à remplir)

            La Puissance à facturer selon le projet est la puissance à facturer actuelle plus la puissance des bornes à installer.
            Le Coût mensuel estimatif (puissance) est calculé comme suit :
            Pour le tarif BR : Si la puissance est supérieure à 50 kW, on enlève 50 kW de la puissance et on multiplie par 0.003 * 31 * 24 * 0.23.
            Pour le tarif M : On multiplie la puissance par 16.14$ * 31 / 30.
            Pour le tarif G9 : On multiplie la puissance par 4.68$ * 31 / 30.
            Pour le tarif G : On enlève 50 kW de la puissance et multiplie le reste par 19.53$.
            Dans votre reponse, donner juste le resultat final mais les formules de calcul
                        
            Mettre cette phrase s’il y’a appel de puissance
            En raison de votre appel de puissance, Nous constatons rapidement que votre cout de la puissance facturee est de (a remplir : donner le montant de la facture mois et année). Les bonnes pratiques en gestion de l'énergie sont de minimiser autant que possible cet appel de puissance. 
            
            Ajouter le paragraphe suivant si la Quantité et puissance des bornes déjà installées est connue 
            Quel est l'impact des bornes actuelles sur ma facture ?
            Les bornes actuelles, représente une puissance ( à remplir), sans l'utilisation d'une gestion adéquate, pourraient très probablement entraîner un appel de puissance pouvant atteindre jusqu'à (a remplir) pour la période analysée (estimation annuelle :  a remplir)
            
            Quel est l'impact et combien peut représenter le cout avec l'installation de mon projet futur ?
            L'ajout de nouvelles bornes, représente une puissance de (a remplir), sans une gestion efficace, votre projet pourrait considérablement augmenter le coût de l'appel de puissance, atteignant potentiellement un montant de (a remplir) par mois, jusqu’à pour un total de 
            (a remplir) annuellement, résultant d'une mauvaise gestion des bornes de recharge.
            
            Comment notre étude peut vous aider ?
            Notre analyse vous permettra de saisir pleinement votre profil de consommation énergétique. Cette approche nous permet d'identifier les moments clés de la journée qui ont un impact significatif sur vos coûts énergétiques. Dans cette perspective, nous vous recommandons vivement d'entreprendre une étude approfondie de votre gestion énergétique. En adoptant les meilleures pratiques et en comprenant vos enjeux spécifiques, vous pourrez réduire de manière substantielle vos dépenses énergétiques. Cette démarche est en parfait alignement avec le plan stratégique 2022-2026 d'Hydro-Québec, axé sur l'optimisation de la consommation d'énergie (Orientation #4). 
            Plus que jamais, la population aura elle aussi un rôle important à jouer dans l’optimisation de la consommation d’énergie. En joignant nos efforts à ceux de l’ensemble des Québécois et Québécoises pour mettre en œuvre des solutions innovantes adaptées à la réalité d’ici, nous pourrons créer un maximum de valeur pour la collectivité.

            
            \n\n{text}\n\n{scenario_text}"""


    # Appeler l'API GPT avec le prompt
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Tu es un assistant pour remplir un rapport d'analyse basé sur les informations suivantes."},
            {"role": "user", "content": full_prompt}
        ]
    )

    # Récupérer la réponse de GPT
    gpt_response = response.choices[0].message.content

    # Générer un fichier Word avec la réponse
    word_file = generate_word_document(gpt_response)
    return word_file

# Interface Gradio
with gr.Blocks(theme='earneleh/paris') as app:
    with gr.Tabs():
        with gr.TabItem("Accueil"):
            gr.Markdown("# Bienvenue dans votre application d'analyse de facture d'hydro Quebec")
            pdf_input = gr.File(label="Charger les fichiers PDF", file_count="multiple", file_types=["pdf"])
            scenario_input = gr.Textbox(label="Entrez le scénario")
            output = gr.File(label="Fichier Word généré")
            analyze_button = gr.Button("Analyser et générer le rapport")
            analyze_button.click(analyze_and_generate, inputs=[pdf_input, scenario_input], outputs=[output])

app.launch(share=True)
